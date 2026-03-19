"""
LectureChor - AI Note Taker
You skip class. We take notes.
Built with Streamlit + Whisper + Claude
"""

import streamlit as st
import openai
import anthropic
import tempfile
import os
import io
import time
import math
import subprocess
import json
import hashlib
import re
from datetime import datetime, timezone, timedelta

# Indian Standard Time (UTC+5:30)
IST = timezone(timedelta(hours=5, minutes=30))

# --- Notes History Storage (session_state backed) ---
# Notes live in session_state so they survive Streamlit reruns within a session.
# They will NOT persist across app restarts on Streamlit Cloud (ephemeral filesystem).
# Users should always download their notes — the sidebar is a convenience, not permanent storage.

def initnotes_store():
    """Initialize notes storage in session_state if not present."""
    if "notes_index" not in st.session_state:
        st.session_state.notes_index = []
    if "notes_data" not in st.session_state:
        st.session_state.notes_data = {}

def load_notes_index():
    initnotes_store()
    return st.session_state.notes_index

def save_notes_index(index):
    st.session_state.notes_index = index

def save_note_entry(title, subject, notes, transcript, word_count, cost, duration, mode, tags=None):
    initnotes_store()
    entry_id = hashlib.md5(f"{title}{datetime.now(IST).isoformat()}".encode()).hexdigest()[:12]
    entry = {
        "id": entry_id,
        "title": title,
        "subject": subject,
        "date": datetime.now(IST).strftime("%B %d, %Y"),
        "time": datetime.now(IST).strftime("%I:%M %p"),
        "word_count": word_count,
        "cost": cost,
        "duration": duration,
        "mode": mode,
        "tags": tags or ([subject] if subject else []),
    }
    st.session_state.notes_data[entry_id] = {"notes": notes, "transcript": transcript}
    st.session_state.notes_index.insert(0, entry)
    return entry_id

def load_note_data(entry_id):
    initnotes_store()
    return st.session_state.notes_data.get(entry_id, None)

def delete_note_entry(entry_id):
    initnotes_store()
    st.session_state.notes_index = [e for e in st.session_state.notes_index if e["id"] != entry_id]
    st.session_state.notes_data.pop(entry_id, None)

# --- Page Config ---
st.set_page_config(
    page_title="LectureChor",
    page_icon="🎓",
    layout="centered",
    initial_sidebar_state="collapsed"
)

# --- Custom CSS ---
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800&family=Space+Grotesk:wght@400;500;600;700&display=swap');

.stApp { font-family: 'Inter', -apple-system, BlinkMacSystemFont, sans-serif; background: #0B0F1A; color: #E2E8F0; }
.stApp::before { content: ''; position: fixed; top: 0; left: 0; width: 100%; height: 100%; background: radial-gradient(ellipse at 20% 50%, rgba(99,102,241,0.08) 0%, transparent 50%), radial-gradient(ellipse at 80% 20%, rgba(168,85,247,0.06) 0%, transparent 50%), radial-gradient(ellipse at 50% 80%, rgba(59,130,246,0.05) 0%, transparent 50%); pointer-events: none; z-index: 0; }

.hero { text-align: center; padding: 2rem 0 1rem; position: relative; }
.hero-icon { font-size: 3.5rem; margin-bottom: 0.5rem; display: block; filter: drop-shadow(0 0 20px rgba(99,102,241,0.4)); }
.hero h1 { font-family: 'Space Grotesk', sans-serif; font-size: 2.8rem; font-weight: 700; background: linear-gradient(135deg, #818CF8 0%, #C084FC 40%, #F472B6 70%, #FB923C 100%); -webkit-background-clip: text; -webkit-text-fill-color: transparent; margin-bottom: 0.3rem; letter-spacing: -1.5px; }
.hero .tagline { font-size: 1.1rem; color: #94A3B8; font-weight: 400; font-style: italic; }
.hero .sub-tagline { font-size: 0.8rem; color: #475569; margin-top: 0.3rem; }

.glass-card { background: rgba(15,23,42,0.6); border: 1px solid rgba(99,102,241,0.15); border-radius: 20px; padding: 2rem; backdrop-filter: blur(20px); box-shadow: 0 8px 32px rgba(0,0,0,0.3), inset 0 1px 0 rgba(255,255,255,0.05); position: relative; overflow: hidden; }

[data-testid="stFileUploader"] { border: 2px dashed rgba(99,102,241,0.3) !important; border-radius: 16px !important; padding: 1.5rem !important; background: rgba(99,102,241,0.03) !important; transition: all 0.3s ease; }
[data-testid="stFileUploader"]:hover { border-color: rgba(99,102,241,0.6) !important; background: rgba(99,102,241,0.06) !important; }

.stButton > button[kind="primary"] { background: linear-gradient(135deg, #6366F1 0%, #8B5CF6 50%, #A855F7 100%) !important; border: none !important; border-radius: 14px !important; padding: 0.8rem 2rem !important; font-weight: 700 !important; font-size: 1rem !important; letter-spacing: 0.5px; transition: all 0.3s ease !important; box-shadow: 0 4px 20px rgba(99,102,241,0.4) !important; text-transform: uppercase; }
.stButton > button[kind="primary"]:hover { transform: translateY(-3px) scale(1.02) !important; box-shadow: 0 8px 30px rgba(99,102,241,0.5) !important; }

.status-card { padding: 1rem 1.4rem; border-radius: 14px; margin: 0.5rem 0; font-size: 0.95rem; backdrop-filter: blur(10px); }
.status-processing { background: rgba(251,191,36,0.08); border-left: 4px solid #F59E0B; color: #FCD34D; }
.status-done { background: rgba(16,185,129,0.08); border-left: 4px solid #10B981; color: #6EE7B7; }
.status-error { background: rgba(239,68,68,0.08); border-left: 4px solid #EF4444; color: #FCA5A5; }

.stats-row { display: flex; gap: 1rem; margin: 1rem 0; }
.stat-box { flex: 1; background: rgba(15,23,42,0.6); border: 1px solid rgba(99,102,241,0.12); border-radius: 16px; padding: 1.2rem; text-align: center; transition: all 0.3s ease; position: relative; overflow: hidden; }
.stat-box::before { content: ''; position: absolute; top: 0; left: 0; right: 0; height: 2px; background: linear-gradient(90deg, #6366F1, #A855F7); opacity: 0; transition: opacity 0.3s ease; }
.stat-box:hover { transform: translateY(-3px); border-color: rgba(99,102,241,0.3); box-shadow: 0 8px 25px rgba(99,102,241,0.15); }
.stat-box:hover::before { opacity: 1; }
.stat-box .num { font-family: 'Space Grotesk', sans-serif; font-size: 1.6rem; font-weight: 700; background: linear-gradient(135deg, #818CF8, #C084FC); -webkit-background-clip: text; -webkit-text-fill-color: transparent; }
.stat-box .label { font-size: 0.7rem; color: #64748B; text-transform: uppercase; letter-spacing: 1.5px; font-weight: 600; margin-top: 0.4rem; }

.password-screen { max-width: 400px; margin: 4rem auto; text-align: center; }
.password-screen .lock-icon { font-size: 3rem; margin-bottom: 1rem; display: block; }
.password-screen h2 { font-family: 'Space Grotesk', sans-serif; font-size: 1.8rem; font-weight: 700; background: linear-gradient(135deg, #818CF8, #C084FC); -webkit-background-clip: text; -webkit-text-fill-color: transparent; margin-bottom: 0.3rem; }
.password-screen .subtitle { color: #64748B; font-size: 0.9rem; margin-bottom: 1.5rem; }

.stDownloadButton > button { width: 100%; border-radius: 12px !important; padding: 0.7rem !important; font-weight: 600 !important; background: rgba(99,102,241,0.08) !important; border: 1px solid rgba(99,102,241,0.2) !important; color: #C7D2FE !important; transition: all 0.3s ease !important; }
.stDownloadButton > button:hover { border-color: rgba(99,102,241,0.5) !important; background: rgba(99,102,241,0.15) !important; transform: translateY(-2px) !important; }

.section-divider { display: flex; align-items: center; gap: 1rem; margin: 1.5rem 0; }
.section-divider .line { flex: 1; height: 1px; background: linear-gradient(90deg, transparent, rgba(99,102,241,0.2), transparent); }
.section-divider .icon { color: #6366F1; font-size: 1.2rem; }

.feature-pills { display: flex; justify-content: center; gap: 0.8rem; margin: 1rem 0; flex-wrap: wrap; }
.pill { background: rgba(99,102,241,0.08); border: 1px solid rgba(99,102,241,0.15); border-radius: 100px; padding: 0.4rem 1rem; font-size: 0.75rem; color: #A5B4FC; font-weight: 500; }

header[data-testid="stHeader"] { background: transparent; }
.block-container { max-width: 780px; padding-top: 1rem; }
hr { border-color: rgba(99,102,241,0.1) !important; }
.stTextInput input { border-radius: 12px !important; border: 1px solid rgba(99,102,241,0.2) !important; background: rgba(15,23,42,0.8) !important; color: #E2E8F0 !important; }
.stTextInput input:focus { border-color: #6366F1 !important; box-shadow: 0 0 0 3px rgba(99,102,241,0.15) !important; }
.stSelectbox > div > div { border-radius: 12px !important; border: 1px solid rgba(99,102,241,0.2) !important; background: rgba(15,23,42,0.8) !important; }

.app-footer { text-align: center; padding: 2rem 0 1rem; margin-top: 3rem; position: relative; }
.app-footer::before { content: ''; position: absolute; top: 0; left: 20%; right: 20%; height: 1px; background: linear-gradient(90deg, transparent, rgba(99,102,241,0.2), transparent); }
.app-footer .footer-brand { font-family: 'Space Grotesk', sans-serif; font-size: 0.85rem; color: #475569; font-weight: 500; }
.app-footer .footer-sub { font-size: 0.7rem; color: #334155; margin-top: 0.3rem; }
.app-footer .easter-egg { font-size: 0.65rem; color: #1E293B; margin-top: 0.5rem; cursor: default; transition: color 0.5s ease; }
.app-footer .easter-egg:hover { color: #6366F1; }

.note-card { background: rgba(15,23,42,0.6); border: 1px solid rgba(99,102,241,0.12); border-radius: 14px; padding: 1rem 1.2rem; margin: 0.5rem 0; transition: all 0.3s ease; }
.note-card:hover { border-color: rgba(99,102,241,0.3); box-shadow: 0 4px 15px rgba(99,102,241,0.1); }
.note-card .note-title { font-family: 'Space Grotesk', sans-serif; font-weight: 600; color: #C7D2FE; font-size: 0.95rem; }
.note-card .note-meta { font-size: 0.75rem; color: #64748B; margin-top: 0.3rem; }

.mode-badge { display: inline-block; padding: 0.15rem 0.6rem; border-radius: 100px; font-size: 0.65rem; font-weight: 600; text-transform: uppercase; letter-spacing: 0.5px; }
.mode-pickpocket { background: rgba(251,191,36,0.15); color: #FCD34D; border: 1px solid rgba(251,191,36,0.3); }
.mode-robbery { background: rgba(99,102,241,0.15); color: #A5B4FC; border: 1px solid rgba(99,102,241,0.3); }
.mode-heist { background: rgba(168,85,247,0.15); color: #D8B4FE; border: 1px solid rgba(168,85,247,0.3); }

.tag-pill { display: inline-block; padding: 0.1rem 0.5rem; border-radius: 100px; font-size: 0.65rem; background: rgba(99,102,241,0.1); color: #818CF8; border: 1px solid rgba(99,102,241,0.15); margin: 0.1rem; }

/* Hide Streamlit branding */
#MainMenu { visibility: hidden; }
footer { visibility: hidden !important; }
footer::after { visibility: hidden !important; content: ''; }
[data-testid="stToolbar"] { display: none !important; }

</style>
""", unsafe_allow_html=True)

# --- Constants ---
MAX_CHUNK_SIZE_MB = 24
WHISPER_MODEL = "whisper-1"
CLAUDE_MODEL = "claude-sonnet-4-20250514"

# --- Note Modes ---
NOTE_MODES = {
    "PickPocket 🤏": {
        "key": "pickpocket",
        "desc": "Quick cheatsheet — key points, case names, sections only",
        "max_tokens": 2500,
        "instruction": "Create a CONCISE cheatsheet-style note. Think exam-eve revision. Include: One-line lecture summary. Bullet-point key concepts (2-3 words each with brief explanation). Case names with one-line ratio. Section/Article numbers with one-line description. 5 most exam-critical points. Keep it SHORT and PUNCHY. No fluff. Every word must earn its place. Follow the lecture flow order. At the end, add a QUICK REFERENCE section grouping all cases and all statutory provisions."
    },
    "Robbery 🔫": {
        "key": "robbery",
        "desc": "Balanced notes — thorough but not exhaustive",
        "max_tokens": 8000,
        "instruction": "Create MEDIUM-LENGTH structured notes. Balance between detail and brevity. Include: Lecture overview (3-4 sentences). Key topics in lecture flow order with clear explanations and examples from the professor. Case law with ratio and how the professor discussed it. Important definitions and statutory provisions. Exam-relevant points. Professor's key insights. Follow the EXACT flow of the lecture. After the main lecture-flow notes, add reference sections: CASE LAW COMPENDIUM with all cases consolidated, STATUTORY PROVISIONS with all sections/articles, KEY DEFINITIONS with all legal terms."
    },
    "Heist 🏴‍☠️": {
        "key": "heist",
        "desc": "Full detailed notes — everything captured in depth",
        "max_tokens": 16000,
        "instruction": "Create EXHAUSTIVE, DETAILED notes capturing EVERYTHING. This is the full heist. PART 1 - LECTURE FLOW NOTES (follow the exact order of the lecture): For each topic as it appears in the lecture provide detailed explanation, the professor's exact analysis and reasoning, every case mentioned with full discussion as the professor presented it, statutory provisions in context, examples and analogies used, cross-references the professor made. DO NOT shorten or summarize the professor's discussion of cases - capture it fully. PART 2 - REFERENCE APPENDIX: 1. CASE LAW COMPENDIUM - Every case mentioned with case name, court and year, key ratio decidendi, how the professor contextualized it, exam tip. 2. KEY DEFINITIONS AND LEGAL TERMS. 3. STATUTORY PROVISIONS - All Acts, Sections, Articles with full references. 4. EXAM-RELEVANT POINTS with answer strategies. 5. CONNECTIONS AND CROSS-REFERENCES. 6. PROFESSOR JEPH'S NOTE - 2-3 personal observations and tips."
    }
}

# --- Password Protection ---
APP_PASSWORD = st.secrets.get("APP_PASSWORD", None)

if not APP_PASSWORD:
    # Local development fallback — will show warning
    APP_PASSWORD = "lecturechor2026"
    st.warning("⚠️ No APP_PASSWORD set in Secrets. Using default password. Set a proper one before sharing with classmates!")

def check_password():
    if "authenticated" not in st.session_state:
        st.session_state.authenticated = False
    if st.session_state.authenticated:
        return True

    st.markdown("""
    <div class="password-screen">
        <span class="lock-icon">🔒</span>
        <h2>LectureChor</h2>
        <p class="subtitle">You skip class. We take notes.</p>
    </div>
    """, unsafe_allow_html=True)

    pwd = st.text_input("Password", type="password", key="pwd_input",
                         label_visibility="collapsed", placeholder="Enter password to continue...")

    if st.button("Unlock", use_container_width=True, type="primary"):
        if pwd == APP_PASSWORD:
            st.session_state.authenticated = True
            st.rerun()
        else:
            st.error("Wrong password. Are you sure you're a LectureChor?")

    st.markdown("""
    <div style="text-align: center; margin-top: 2rem;">
        <span style="font-size: 0.75rem; color: #334155;">Hint: Ask a fellow chor</span>
    </div>
    """, unsafe_allow_html=True)
    return False

def get_file_size_mb(filepath):
    return os.path.getsize(filepath) / (1024 * 1024)

def compute_audio_hash(filepath):
    h = hashlib.md5()
    with open(filepath, "rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            h.update(chunk)
    return h.hexdigest()

def split_audio(filepath):
    size_mb = get_file_size_mb(filepath)
    if size_mb <= MAX_CHUNK_SIZE_MB:
        return [filepath]

    result = subprocess.run(
        ["ffprobe", "-v", "quiet", "-show_entries", "format=duration",
         "-of", "csv=p=0", filepath],
        capture_output=True, text=True
    )
    total_duration = float(result.stdout.strip())
    num_chunks = math.ceil(size_mb / MAX_CHUNK_SIZE_MB)
    chunk_duration = total_duration / num_chunks
    OVERLAP_SECONDS = 30  # Overlap to avoid cutting mid-sentence

    chunks = []
    tmp_dir = tempfile.mkdtemp()
    for i in range(num_chunks):
        chunk_path = os.path.join(tmp_dir, f"chunk_{i+1:03d}.mp3")
        # Start earlier for chunks after the first to capture overlap
        actual_start = max(0, i * chunk_duration - OVERLAP_SECONDS) if i > 0 else 0
        actual_duration = chunk_duration + (OVERLAP_SECONDS if i > 0 else 0)
        # Don't overshoot total duration
        if actual_start + actual_duration > total_duration:
            actual_duration = total_duration - actual_start

        cmd = [
            "ffmpeg", "-y", "-i", filepath,
            "-ss", str(actual_start), "-t", str(actual_duration),
            "-acodec", "libmp3lame", "-ab", "64k", "-ar", "16000",
            chunk_path
        ]
        subprocess.run(cmd, capture_output=True, check=True)
        chunks.append(chunk_path)
    return chunks

def transcribe_audio(filepath, subject, progress_callback=None):
    client = openai.OpenAI(api_key=st.secrets["OPENAI_API_KEY"])
    audio_hash = compute_audio_hash(filepath)

    # Cache in session_state instead of filesystem
    cache_key = f"transcript_cache_{audio_hash}"
    if cache_key in st.session_state:
        return st.session_state[cache_key]

    chunks = split_audio(filepath)
    full_transcript = []

    for i, chunk_path in enumerate(chunks):
        if progress_callback:
            progress_callback(f"Transcribing chunk {i+1}/{len(chunks)}..." if len(chunks) > 1 else "Transcribing...")

        with open(chunk_path, "rb") as f:
            prompt_text = "This is a law lecture"
            if subject:
                prompt_text += f" on {subject}"
            prompt_text += ". Legal terms, case names, and some Hindi words may appear."

            response = client.audio.transcriptions.create(
                model=WHISPER_MODEL, file=f, language="en",
                prompt=prompt_text, response_format="text", temperature=0.0
            )
        full_transcript.append(response)

    result = " ".join(full_transcript)
    st.session_state[cache_key] = result
    return result

def extract_case_names(text):
    case_patterns = [
        # Italicized case names: Maneka Gandhi v. Union of India
        r'\*([A-Z][a-zA-Z\s\.\,\'\-]+(?:v\.?|vs\.?)\s+[A-Z][a-zA-Z\s\.\,\'\-]+(?:\([^)]*\))?)\*',
        # Standard format: State of X v. Person Name (Year) or with "and" etc.
        r'(?<!\w)([A-Z][a-zA-Z\s\.\,\'\-]+(?:v\.?|vs\.?)\s+[A-Z][a-zA-Z\s\.\,\'\-]+(?:\(\d{4}\))?)',
        # Linked case names: [Case Name](url)
        r'\[([A-Z][a-zA-Z\s\.\,\'\-]+(?:v\.?|vs\.?)\s+[A-Z][a-zA-Z\s\.\,\'\-]+(?:\([^)]*\))?)\]\(',
    ]

    # Common false-positive phrases to exclude
    false_positives = {"Article vs Section", "State vs Central", "Law vs Justice", "Rights vs Duties"}

    cases = set()
    for pattern in case_patterns:
        matches = re.findall(pattern, text)
        for match in matches:
            clean = match.strip().strip('*').strip()
            # Must be reasonable length, contain "v." or "vs.", and not be a false positive
            if 10 < len(clean) < 120 and clean not in false_positives:
                # Trim trailing whitespace and punctuation
                clean = re.sub(r'[\s,\.]+$', '', clean)
                cases.add(clean)
    return list(cases)

def generate_case_links(cases):
    links = {}
    for case in cases:
        search_query = case.replace(" ", "+")
        links[case] = f"https://indiankanoon.org/search/?formInput={search_query}"
    return links

def generate_notes(transcript, subject, filename, mode_key="heist"):
    client = anthropic.Anthropic(api_key=st.secrets["ANTHROPIC_API_KEY"])
    subject_expertise = subject if subject else "Indian Law"

    mode_config = None
    for name, cfg in NOTE_MODES.items():
        if cfg["key"] == mode_key:
            mode_config = cfg
            break
    if not mode_config:
        mode_config = list(NOTE_MODES.values())[2]

    system_prompt = f"""You are a note-taking AI that converts law lecture transcripts into class notes. Your output must read EXACTLY like notes taken by a sharp LLM (Masters of Law) student typing furiously during a live lecture in an Indian law university. NOT a textbook. NOT a polished summary. Real lecture notes.

SUBJECT EXPERTISE: {subject_expertise}

VOICE AND TONE:
- Write as if you ARE the student sitting in class, capturing the professor's words in real time
- Mix the professor's voice with your own shorthand. "The court says..." "The prof argues..." "So the question is..."
- When the professor gives an opinion, capture it raw like: "Prof thinks teachers should be workmen under the act" or "Sir believes if the doctor's report wasn't there, the case could have gone differently"
- When the professor makes a joke or aside, include it naturally
- Capture the professor's emphasis: "IMP", "very imp for exam", "read this case carefully"

ABBREVIATIONS - USE THESE NATURALLY AND INCONSISTENTLY (like a real person):
bc = because, defn = definition, estb = establishment, diff = different, govt = government, consti = constitutional, wrt = with regard to, eco = economic, mnf = manufacturing, dom = dominant, esp = especially, u/ur = you/your, prof = professor, b/w = between, v = versus, sec = section, art = article, para = paragraph, SC = Supreme Court, HC = High Court

STRUCTURE - FOLLOW THE LECTURE FLOW:
- Follow the EXACT order the professor discusses topics. Do NOT reorganize into neat textbook sections
- When the professor jumps between topics, your notes jump too
- When the professor circles back, write "Coming back to..." or "So..."
- Use headings ONLY for major topic shifts or case names
- Case names get their own line as a heading

CASE LAW - CRITICAL: Cases must be discussed AS THE PROFESSOR DISCUSSED THEM:
- Case name with citation if mentioned
- Facts as the professor narrated them (keep their storytelling style)
- The professor's analysis woven in: "The court gets into..." "The court says..."
- Quote specific paragraph numbers if mentioned: "(para 56)" or "Read page 313"
- Professor's commentary: "This was called a broader interpretation judgment"
- How this case connects to others: "This was overruled in BWSSB"
- Add Indian Kanoon links: [Case Name](https://indiankanoon.org/search/?formInput=CASE+NAME+HERE)
DO NOT format cases as Facts/Issues/Held/Ratio boxes. That's textbook style, not lecture notes.

WHAT TO CAPTURE:
- Every legal concept, case name, statutory provision, professor's opinion
- Practical examples and analogies ("Is Amazon an industry?" "Is a beauty parlour an industry?")
- Questions the professor asks the class and answers
- Reading references, cross-references, Hindi words with context

FORMATTING:
- Markdown headings sparingly - only for new modules, major topics, case names
- Bold key legal terms on FIRST mention only
- Short bullet fragments for lists of elements/tests
- Keep paragraphs as flowing lecture capture, not neat bullet summaries
- Note unclear parts as [Unclear in recording]
- Use dashes freely, use fragments naturally
- Do NOT over-polish. Keep it raw like real notes.

DO NOT:
- Write like a textbook or encyclopedia
- Create neat Facts/Issues/Held/Ratio boxes
- Add sections the professor didn't discuss
- Use formal language like "It is pertinent to note" or "The Hon'ble Court observed"
- Over-structure with Roman numerals and nested sub-headings
- Make every sentence a bullet point

CRITICAL INSTRUCTION: {mode_config['instruction']}

CASE LAW LINKING: For EVERY case you mention, add a link: [Case Name](https://indiankanoon.org/search/?formInput=CASE+NAME+ENCODED)"""

    subject_line = f"**Subject: {subject}**\n" if subject else ""
    message = client.messages.create(
        model=CLAUDE_MODEL,
        max_tokens=mode_config["max_tokens"],
        system=system_prompt,
        messages=[{
            "role": "user",
            "content": f"""{subject_line}**Source:** {filename}
Date: {datetime.now(IST).strftime('%B %d, %Y')}

Below is the full transcript of a law lecture. Create comprehensive, structured notes following your format. Be exhaustive. Do NOT shorten or omit the professor's discussion of any case law.

---

TRANSCRIPT:
{transcript}

---

Generate complete structured notes now, Professor Jeph."""
        }]
    )

    notes_text = message.content[0].text

    # Detect if Claude hit the token limit and notes were cut off
    if message.stop_reason == "max_tokens":
        notes_text += "\n\n---\n⚠️ Note: These notes were truncated because the lecture was very long. Consider using Robbery or PickPocket mode for lengthy lectures, or splitting the audio into smaller files."

    return notes_text

def reformat_uploaded_notes(doc_text, subject, mode_key="robbery"):
    client = anthropic.Anthropic(api_key=st.secrets["ANTHROPIC_API_KEY"])
    subject_expertise = subject if subject else "Indian Law"

    mode_config = None
    for name, cfg in NOTE_MODES.items():
        if cfg["key"] == mode_key:
            mode_config = cfg
            break
    if not mode_config:
        mode_config = list(NOTE_MODES.values())[1]

    system_prompt = f"""You are a note-taking AI that restructures existing law notes into the style of a sharp LLM student's lecture notes from an Indian law university. NOT a textbook. NOT a polished summary. Real lecture notes style.

You have been given existing notes on {subject_expertise} that need to be reorganized and enhanced.

Your task: 1. Restructure these notes following lecture-capture style. 2. Add exam tips and cross-references where relevant. 3. Link all case names to Indian Kanoon: [Case Name](https://indiankanoon.org/search/?formInput=CASE+NAME+ENCODED). 4. Preserve ALL original content. 5. Enhance definitions, add context to cases.

VOICE AND TONE:
- Write as if a sharp student captured these in class
- Use shorthand naturally: bc, defn, estb, diff, govt, consti, wrt, eco, dom, esp, sec, art, para, SC, HC
- Mix abbreviations inconsistently like a real person
- Keep professor opinions and commentary raw

CASE LAW - discuss cases narratively, NOT as Facts/Issues/Held/Ratio boxes. Use storytelling style.

FORMATTING:
- Markdown headings sparingly - only for major topics and case names
- Bold key terms on first mention only
- Short bullet fragments for tests and elements
- Flowing paragraphs for discussion, not wall of bullets
- Note unclear parts as [Unclear in recording]

{mode_config['instruction']}"""

    message = client.messages.create(
        model=CLAUDE_MODEL,
        max_tokens=mode_config["max_tokens"],
        system=system_prompt,
        messages=[{
            "role": "user",
            "content": f"Please reorganize and enhance these notes on {subject_expertise}:\n\n---\n{doc_text}\n---\n\nApply your signature LectureChor format."
        }]
    )

    notes_text = message.content[0].text
    if message.stop_reason == "max_tokens":
        notes_text += "\n\n---\n⚠️ Note: These reformatted notes were truncated due to length. Try using PickPocket mode for very long documents."

    return notes_text

def create_docx_bytes(notes_text, subject, filename, include_case_links=True):
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.15

    for level, (size, color) in enumerate([
        (Pt(18), RGBColor(0x1B, 0x3A, 0x5C)),
        (Pt(14), RGBColor(0x2E, 0x75, 0xB6)),
        (Pt(12), RGBColor(0x40, 0x40, 0x40)),
    ], 1):
        h = doc.styles[f'Heading {level}']
        h.font.name = 'Calibri'
        h.font.size = size
        h.font.color.rgb = color
        h.font.bold = True

    doc.add_paragraph()
    doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("LECTURE NOTES")
    r.font.size = Pt(32)
    r.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
    r.font.bold = True

    if subject:
        s = doc.add_paragraph()
        s.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = s.add_run(subject.upper())
        r.font.size = Pt(18)
        r.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

    m = doc.add_paragraph()
    m.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = m.add_run(f"Generated: {datetime.now(IST).strftime('%B %d, %Y')}\nSource: {filename}\nBy LectureChor AI")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.add_page_break()

    for line in notes_text.split('\n'):
        line = line.rstrip()
        if not line:
            continue
        if line.startswith('### '):
            doc.add_heading(line[4:].strip().replace('**', ''), level=3)
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip().replace('**', ''), level=2)
        elif line.startswith('# '):
            doc.add_heading(line[2:].strip().replace('**', ''), level=1)
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            addformatted_runs(p, line[2:].strip())
        elif line.startswith('  - ') or line.startswith('  * '):
            p = doc.add_paragraph(style='List Bullet 2')
            addformatted_runs(p, line[4:].strip())
        elif line.startswith('---'):
            continue
        else:
            p = doc.add_paragraph()
            addformatted_runs(p, line)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()

def create_transcript_docx(transcript, subject, filename):
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.5

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("LECTURE TRANSCRIPT")
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
    r.font.bold = True

    if subject:
        s = doc.add_paragraph()
        s.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = s.add_run(subject)
        r.font.size = Pt(14)
        r.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

    m = doc.add_paragraph()
    m.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = m.add_run(f"Source: {filename}\nDate: {datetime.now(IST).strftime('%B %d, %Y')}\nWord Count: {len(transcript.split()):,}")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.add_paragraph()
    words = transcript.split()
    for i in range(0, len(words), 200):
        doc.add_paragraph(" ".join(words[i:i+200]))

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()

def addformatted_runs(paragraph, text):
    parts = re.split(r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if part.startswith('***') and part.endswith('***'):
            r = paragraph.add_run(part[3:-3])
            r.bold = True
            r.italic = True
        elif part.startswith('**') and part.endswith('**'):
            r = paragraph.add_run(part[2:-2])
            r.bold = True
        elif part.startswith('*') and part.endswith('*'):
            r = paragraph.add_run(part[1:-1])
            r.italic = True
        else:
            paragraph.add_run(part)

def extract_text_from_docx(uploaded_file):
    from docx import Document
    doc = Document(io.BytesIO(uploaded_file.getvalue()))
    full_text = []
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)
    return "\n".join(full_text)

# ======= MAIN APP =======
if not check_password():
    st.stop()

# --- Sidebar: My Notes ---
with st.sidebar:
    st.markdown("### 📚 My Notes")
    st.markdown("---")

    notes_index = load_notes_index()

    all_tags = set()
    for entry in notes_index:
        for tag in entry.get("tags", []):
            all_tags.add(tag)

    if all_tags:
        filter_tag = st.selectbox("Filter by subject", ["All"] + sorted(all_tags), key="filter_tag")
    else:
        filter_tag = "All"

    filtered_notes = notes_index if filter_tag == "All" else [
        e for e in notes_index if filter_tag in e.get("tags", [])
    ]

    if not filtered_notes:
        st.markdown("<p style='color: #64748B; font-size: 0.85rem; text-align: center; padding: 2rem 0;'>No notes yet. Generate your first one!</p>", unsafe_allow_html=True)

    else:
        for entry in filtered_notes:
            mode_class = f"mode-{entry.get('mode', 'robbery')}"
            mode_label = entry.get("mode", "robbery").upper()
            st.markdown(f"""
            <div class="note-card">
                <div class="note-title">{entry['title']}</div>
                <div class="note-meta">{entry['date']} at {entry['time']} &nbsp; <span class="mode-badge {mode_class}">{mode_label}</span></div>
            </div>
            """, unsafe_allow_html=True)

            if st.button(f"📖 Open", key=f"open_{entry['id']}", use_container_width=True):
                st.session_state['viewing_note_id'] = entry['id']
                st.session_state['viewing_note_entry'] = entry

    st.markdown("---")
    st.markdown("### 🔄 Organise Your Notes")
    uploaded_doc = st.file_uploader("Upload a Word document to reformat", type=["docx"],
        key="reformat_uploader", help="Upload your own notes and Prof Jeph will restructure them")

    if uploaded_doc:
        reformat_subject = st.text_input("Subject for these notes", key="reformat_subject", placeholder="e.g., Criminal Law")
        reformat_mode = st.selectbox("Note style", list(NOTE_MODES.keys()), index=1, key="reformat_mode")
        reformat_tags_input = st.text_input("Tags (comma separated)", key="reformat_tags", placeholder="e.g., Criminal, IPC, Semester 3")

        if st.button("🔄 Reformat Notes", use_container_width=True, type="primary", key="reformat_btn"):
            with st.status("🧠 Prof. Jeph is reorganizing your notes...", expanded=True) as status:
                doc_text = extract_text_from_docx(uploaded_doc)
                st.write(f"Extracted {len(doc_text.split()):,} words from document...")
                start_time = time.time()
                mode_key = NOTE_MODES[reformat_mode]["key"]
                reformatted = reformat_uploaded_notes(doc_text, reformat_subject, mode_key)
                reformat_time = time.time() - start_time
                st.write(f"✅ Reformatted in {reformat_time:.0f}s")
                status.update(label="✅ Notes reorganized!", state="complete")

            word_count = len(doc_text.split())
            # Reformat = Claude only (no Whisper)
            INR_PER_USD = 85
            claude_input_cost = (word_count * 1.3 / 1_000_000) * 3 * INR_PER_USD
            reformat_mode_cfg = NOTE_MODES[reformat_mode]
            claude_output_cost = (reformat_mode_cfg["max_tokens"] / 1_000_000) * 15 * INR_PER_USD
            display_cost = max(1, int(claude_input_cost + claude_output_cost))

            tags = [t.strip() for t in reformat_tags_input.split(",") if t.strip()] if reformat_tags_input else []
            if reformat_subject and reformat_subject not in tags:
                tags.insert(0, reformat_subject)

            save_note_entry(title=uploaded_doc.name, subject=reformat_subject, notes=reformatted,
                transcript=doc_text, word_count=word_count, cost=display_cost,
                duration=int(reformat_time), mode=mode_key, tags=tags)

            st.markdown("### ✅ Reformatted Notes")
            st.markdown(reformatted)
            st.rerun()

# Hero
st.markdown("""
<div class="hero">
    <span class="hero-icon">🎓</span>
    <h1>LectureChor</h1>
    <p class="tagline">You skip class. We take notes.</p>
    <p class="sub-tagline">Powered by KJ who actually paid attention in class</p>
</div>
""", unsafe_allow_html=True)

st.markdown("""
<div class="feature-pills">
    <span class="pill">🎙️ Whisper Transcription</span>
    <span class="pill">🧠 Claude AI Notes</span>
    <span class="pill">📄 Word Export</span>
    <span class="pill">🇮🇳 Hindi + English</span>
    <span class="pill">🔗 Case Law Links</span>
</div>
""", unsafe_allow_html=True)

st.markdown("""
<div class="section-divider">
    <div class="line"></div>
    <span class="icon">✨</span>
    <div class="line"></div>
</div>
""", unsafe_allow_html=True)

# Upload section
col1, col2 = st.columns([2, 1])
with col1:
    uploaded_file = st.file_uploader("Drop your lecture audio here",
        type=["mp3", "m4a", "wav", "mp4", "ogg", "flac", "webm"],
        help="Supports MP3, M4A, WAV, MP4, OGG, FLAC, WebM")
with col2:
    subject = st.text_input("Subject", placeholder="e.g., Constitutional Law",
        help="Helps AI focus on the right area of law")
    note_mode = st.selectbox("Note Style", list(NOTE_MODES.keys()), index=2,
        help="PickPocket = cheatsheet, Robbery = balanced, Heist = full detail")

selected_mode = NOTE_MODES[note_mode]
st.markdown(f"""
<div style="text-align: center; margin: 0.5rem 0;">
    <span class="mode-badge mode-{selected_mode['key']}">{selected_mode['key'].upper()}</span>
    <span style="color: #64748B; font-size: 0.8rem; margin-left: 0.5rem;">{selected_mode['desc']}</span>
</div>
""", unsafe_allow_html=True)

# Process button
if uploaded_file is not None:
    file_size_mb = len(uploaded_file.getvalue()) / (1024 * 1024)
    st.markdown(f"""
    <div class="stats-row">
        <div class="stat-box"><div class="num">{uploaded_file.name.split('.')[-1].upper()}</div><div class="label">Format</div></div>
        <div class="stat-box"><div class="num">{file_size_mb:.1f} MB</div><div class="label">File size</div></div>
        <div class="stat-box"><div class="num">{subject or 'Auto'}</div><div class="label">Subject</div></div>
    </div>
    """, unsafe_allow_html=True)

    if st.button("⚡ Generate Notes", use_container_width=True, type="primary"):
        with tempfile.NamedTemporaryFile(delete=False, suffix=f".{uploaded_file.name.split('.')[-1]}") as tmp:
            tmp.write(uploaded_file.getvalue())
            tmp_path = tmp.name

        try:
            with st.status("🎙️ Transcribing audio with Whisper AI...", expanded=True) as status:
                st.write(f"Processing {file_size_mb:.1f}MB audio file...")
                start_time = time.time()
                transcript = transcribe_audio(tmp_path, subject)
                word_count = len(transcript.split())
                transcribe_time = time.time() - start_time
                st.write(f"✅ Transcription complete — {word_count:,} words in {transcribe_time:.0f}s")
                status.update(label="✅ Transcription complete", state="complete")

            mode_key = selected_mode["key"]
            with st.status("🧠 Professor Jeph is preparing your notes...", expanded=True) as status:
                st.write("Analyzing transcript for legal concepts, case law, definitions...")
                start_time = time.time()
                notes = generate_notes(transcript, subject, uploaded_file.name, mode_key)
                notes_time = time.time() - start_time
                st.write(f"✅ Notes generated in {notes_time:.0f}s")
                status.update(label="✅ Notes generated by Prof. Jeph", state="complete")

            total_time = transcribe_time + notes_time
            # Realistic cost: Whisper ($0.006/min) + Claude Sonnet input ($3/M tokens) + output ($15/M tokens)
            INR_PER_USD = 85
            est_audio_minutes = word_count / 150  # ~150 words/min in speech
            whisper_cost = est_audio_minutes * 0.006 * INR_PER_USD
            claude_input_cost = (word_count * 1.3 / 1_000_000) * 3 * INR_PER_USD
            claude_output_cost = (selected_mode["max_tokens"] / 1_000_000) * 15 * INR_PER_USD
            display_cost = max(1, int(whisper_cost + claude_input_cost + claude_output_cost))

            st.session_state['last_notes'] = notes
            st.session_state['last_transcript'] = transcript
            st.session_state['last_word_count'] = word_count
            st.session_state['last_total_time'] = total_time
            st.session_state['last_cost'] = display_cost
            st.session_state['last_filename'] = uploaded_file.name
            st.session_state['last_subject'] = subject
            st.session_state['last_mode'] = mode_key

            tags = [subject] if subject else []
            save_note_entry(title=uploaded_file.name, subject=subject, notes=notes,
                transcript=transcript, word_count=word_count, cost=display_cost,
                duration=int(total_time), mode=mode_key, tags=tags)

        except Exception as e:
            st.markdown(f"""<div class="status-card status-error">❌ <strong>Error:</strong> {str(e)}</div>""", unsafe_allow_html=True)
            st.error(f"Full error: {e}")

        finally:
            os.unlink(tmp_path)

# Display results from session state (persists across reruns)
if 'last_notes' in st.session_state:
    notes = st.session_state['last_notes']
    transcript = st.session_state['last_transcript']
    word_count = st.session_state.get('last_word_count', 0)
    total_time = st.session_state.get('last_total_time', 0)
    display_cost = st.session_state.get('last_cost', 0)
    filename = st.session_state.get('last_filename', 'lecture')
    subject_saved = st.session_state.get('last_subject', '')
    mode_key = st.session_state.get('last_mode', 'heist')

    st.markdown("""
    <div class="section-divider">
        <div class="line"></div>
        <span class="icon">📚</span>
        <div class="line"></div>
    </div>
    """, unsafe_allow_html=True)

    st.markdown("### 📋 Your Notes Are Ready")

    st.markdown(f"""
    <div class="stats-row">
        <div class="stat-box"><div class="num">{word_count:,}</div><div class="label">Words transcribed</div></div>
        <div class="stat-box"><div class="num">{total_time:.0f}s</div><div class="label">Total time</div></div>
        <div class="stat-box"><div class="num">₹{display_cost:,}</div><div class="label">Est. cost</div></div>
    </div>
    """, unsafe_allow_html=True)

    st.markdown("#### 📥 Download")
    safe_name = subject_saved.replace(' ', '_')[:20] if subject_saved else "Lecture"
    col_a, col_b, col_c = st.columns(3)

    with col_a:
        docx_bytes = create_docx_bytes(notes, subject_saved, filename)
        st.download_button("📄 Notes (Word)", data=docx_bytes,
            file_name=f"LectureChor_{safe_name}_{datetime.now(IST).strftime('%Y%m%d')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True, key="dl_main_notes")

    with col_b:
        transcript_docx = create_transcript_docx(transcript, subject_saved, filename)
        st.download_button("🎙️ Transcript (Word)", data=transcript_docx,
            file_name=f"Transcript_{safe_name}_{datetime.now(IST).strftime('%Y%m%d')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True, key="dl_main_transcript")

    with col_c:
        md_content = f"# Lecture Notes: {subject_saved or filename}\n*Generated by LectureChor: {datetime.now(IST).strftime('%B %d, %Y')}*\n\n{notes}"
        st.download_button("📝 Markdown", data=md_content,
            file_name=f"LectureChor_{safe_name}_{datetime.now(IST).strftime('%Y%m%d')}.md",
            mime="text/markdown", use_container_width=True, key="dl_main_md")

    st.markdown("""
    <div class="section-divider">
        <div class="line"></div>
        <span class="icon">📖</span>
        <div class="line"></div>
    </div>
    """, unsafe_allow_html=True)

    # Case law links
    cases = extract_case_names(notes)
    if cases:
        st.markdown("#### 🔗 Case Law Quick Links")
        case_links = generate_case_links(cases)
        for case_name, link in case_links.items():
            st.markdown(f"- [*{case_name}*]({link})")
        st.markdown("")

    # Notes display
    st.markdown(notes)

    if st.button("🗑️ Clear Results", key="clear_results"):
        for key in ['last_notes', 'last_transcript', 'last_word_count', 'last_total_time',
                     'last_cost', 'last_filename', 'last_subject', 'last_mode']:
            st.session_state.pop(key, None)
        st.rerun()

# --- Display saved note from sidebar ---
if 'viewing_note_id' in st.session_state:
    view_entry = st.session_state.get('viewing_note_entry', {})
    view_data = load_note_data(st.session_state['viewing_note_id'])

    if view_data:
        st.markdown("""
        <div class="section-divider">
            <div class="line"></div>
            <span class="icon">📂</span>
            <div class="line"></div>
        </div>
        """, unsafe_allow_html=True)

        mode_class = f"mode-{view_entry.get('mode', 'robbery')}"
        mode_label = view_entry.get("mode", "robbery").upper()

        st.markdown(f"### 📖 {view_entry.get('title', 'Saved Note')}")
        st.markdown(f"<span class='mode-badge {mode_class}'>{mode_label}</span> &nbsp; {view_entry.get('date', '')} at {view_entry.get('time', '')}", unsafe_allow_html=True)

        st.markdown(f"""
        <div class="stats-row">
            <div class="stat-box"><div class="num">{view_entry.get('word_count', 0):,}</div><div class="label">Words</div></div>
            <div class="stat-box"><div class="num">{view_entry.get('duration', 0)}s</div><div class="label">Time</div></div>
            <div class="stat-box"><div class="num">₹{view_entry.get('cost', 0)}</div><div class="label">Cost</div></div>
        </div>
        """, unsafe_allow_html=True)

        # Download buttons
        safe_name = view_entry.get("subject", "Lecture").replace(' ', '_')[:20]
        eid = view_entry.get("id", "note")
        col_a, col_b, col_c = st.columns(3)

        with col_a:
            docx_bytes = create_docx_bytes(view_data["notes"], view_entry.get("subject", ""), view_entry.get("title", ""))
            st.download_button("📄 Notes (Word)", data=docx_bytes,
                file_name=f"LectureChor_{safe_name}_{eid}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True, key=f"dl_view_notes_{eid}")

        with col_b:
            transcript_docx = create_transcript_docx(
                view_data.get("transcript", "No transcript available"),
                view_entry.get("subject", ""), view_entry.get("title", ""))
            st.download_button("🎙️ Transcript (Word)", data=transcript_docx,
                file_name=f"Transcript_{safe_name}_{eid}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True, key=f"dl_view_trans_{eid}")

        with col_c:
            md_content = f"# {view_entry.get('title', '')}\n*{view_entry.get('date', '')}*\n\n{view_data['notes']}"
            st.download_button("📝 Markdown", data=md_content,
                file_name=f"LectureChor_{safe_name}_{eid}.md",
                mime="text/markdown", use_container_width=True, key=f"dl_view_md_{eid}")

        # Case law links
        saved_cases = extract_case_names(view_data["notes"])
        if saved_cases:
            st.markdown("#### 🔗 Case Law Quick Links")
            saved_links = generate_case_links(saved_cases)
            for case_name, link in saved_links.items():
                st.markdown(f"- [*{case_name}*]({link})")

        # Full notes
        st.markdown(view_data["notes"])

        if st.button("✕ Close", key="close_saved_note"):
            st.session_state.pop('viewing_note_id', None)
            st.session_state.pop('viewing_note_entry', None)
            st.rerun()

# Footer
st.markdown("""
<div class="app-footer">
    <div class="footer-brand">🎓 Built by KJ for LectureChors</div>
    <div class="footer-sub">Made with sleep deprivation and chai</div>
    <div class="easter-egg">v4.1 — KJ was here — jeph.exe has stopped working — now with PickPocket, Robbery & Heist modes</div>
</div>
""", unsafe_allow_html=True)
